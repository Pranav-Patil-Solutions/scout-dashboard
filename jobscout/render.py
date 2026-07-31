#!/usr/bin/env python3
"""Navy/stone CV + cover-letter renderers. Design is byte-for-byte the canonical
build_cv_docx.py look; content comes from profile.py + a per-job spec dict."""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
INK = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x55, 0x55, 0x55)
RULE = "CFD6DE"
FONT = "Calibri"

HERE = os.path.dirname(os.path.abspath(__file__))
HEADSHOT = os.path.join(os.path.dirname(HERE), "headshot-circular.png")  # ~/Desktop/CV-Template/


# ---------- shared docx helpers ----------
def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val)); node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcPr.append(mar)


def no_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}"); el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def add_bottom_rule(paragraph, color=RULE, size=6, space="2"):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), space); bottom.set(qn("w:color"), color)
    pbdr.append(bottom); pPr.append(pbdr)


def add_hyperlink(paragraph, url, text, color="1F3A5F", underline=False, size=9, bold=False):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink"); hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts"); rFonts.set(qn("w:ascii"), FONT); rPr.append(rFonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(size * 2)); rPr.append(sz)
    if bold:
        rPr.append(OxmlElement("w:b"))
    c = OxmlElement("w:color"); c.set(qn("w:val"), color); rPr.append(c)
    if underline:
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; new_run.append(t)
    hyperlink.append(new_run); paragraph._p.append(hyperlink)


def _run(p, text, bold=False, italic=False, size=9, color=INK, font=FONT):
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.color.rgb = color; r.font.name = font
    return r


# ============================================================
#  CV
# ============================================================
def build_cv(profile, job, out_path):
    """profile: the profile module. job: dict spec (may be empty). out_path: .docx path."""
    cv = job.get("cv", {})
    C = profile.CONTACT

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27); section.page_height = Inches(11.69)
    section.top_margin = Inches(0.35); section.bottom_margin = Inches(0.32)
    section.left_margin = Inches(0.4); section.right_margin = Inches(0.4)

    normal = doc.styles["Normal"]
    normal.font.name = FONT; normal.font.size = Pt(9); normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(0); normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    def para(container, space_before=0, space_after=0, align=None):
        p = container.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.0
        if align:
            p.alignment = align
        return p

    def heading(container, text, space_before=6, space_after=3):
        p = para(container, space_before=space_before, space_after=1)
        _run(p, text, bold=True, size=11.5, color=NAVY)
        add_bottom_rule(p)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def bullets(container, items, size=8):
        for item in items:
            p = container.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(1.5)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.left_indent = Inches(0.14)
            for text, bold in item:
                _run(p, text, bold=bold, size=size, color=NAVY if bold else INK)

    def entry(container, dates, loc, role, company, items, company_desc=None):
        meta = para(container, space_before=4, space_after=0.5)
        _run(meta, f"{dates}  •  {loc}", size=7.2, color=MUTED)
        r = para(container, space_after=0)
        _run(r, role, bold=True, size=9.3, color=NAVY)
        c = para(container, space_after=1.5)
        _run(c, company, bold=True, size=8.6, color=INK)
        if company_desc:
            cd = para(container, space_after=1.5)
            _run(cd, company_desc, italic=True, size=7.6, color=MUTED)
        if items:
            bullets(container, items)

    def side_list(container, items, size=8):
        for bold_part, text in items:
            p = para(container, space_after=1)
            if bold_part:
                _run(p, bold_part, bold=True, size=size, color=INK)
            _run(p, text, size=size, color=INK)

    # ---- HEADER ----
    header_tbl = doc.add_table(rows=1, cols=2)
    header_tbl.autofit = False; no_borders(header_tbl)
    header_tbl.columns[0].width = Inches(6.1); header_tbl.columns[1].width = Inches(1.2)
    name_cell, photo_cell = header_tbl.rows[0].cells
    name_cell.width = Inches(6.1); photo_cell.width = Inches(1.2)
    set_cell_margins(name_cell, end=100)

    p = name_cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    _run(p, C["name"], bold=True, size=20, color=NAVY)

    p = para(name_cell, space_after=3)
    _run(p, cv.get("headline", C["headline"]), bold=True, size=8.6, color=NAVY)

    p = para(name_cell, space_after=1)
    _run(p, "✉ ", size=7.6, color=NAVY, bold=True)
    add_hyperlink(p, f"mailto:{C['email']}", C["email"], size=7.6)
    _run(p, "      ✆ ", size=7.6, color=NAVY, bold=True)
    add_hyperlink(p, f"tel:{C['phone_tel']}", C["phone"], size=7.6)
    _run(p, "      ◎ ", size=7.6, color=NAVY, bold=True)
    _run(p, C["location"], size=7.6)

    p = para(name_cell, space_after=0)
    _run(p, "in ", size=7.6, color=NAVY, bold=True)
    add_hyperlink(p, C["linkedin_url"], C["linkedin"], size=7.6)
    _run(p, "      ⌥ ", size=7.6, color=NAVY, bold=True)
    add_hyperlink(p, C["github_url"], C["github"], size=7.6)

    photo_p = photo_cell.paragraphs[0]; photo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(HEADSHOT):
        photo_p.add_run().add_picture(HEADSHOT, width=Inches(0.85), height=Inches(0.85))
    photo_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0); spacer.paragraph_format.space_before = Pt(0)
    _run(spacer, "", size=2)

    # ---- BODY (two columns) ----
    body_tbl = doc.add_table(rows=1, cols=2)
    body_tbl.autofit = False; no_borders(body_tbl)
    body_tbl.columns[0].width = Inches(4.5); body_tbl.columns[1].width = Inches(2.8)
    main, side = body_tbl.rows[0].cells
    main.width = Inches(4.5); side.width = Inches(2.8)
    set_cell_margins(main, end=200); set_cell_margins(side, start=100)
    main.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    side.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    main.paragraphs[0].text = ""; side.paragraphs[0].text = ""

    # MAIN
    heading(main, "Professional Experience", space_before=0)
    for e in profile.EXPERIENCE:
        entry(main, e["dates"], e["loc"], e["role"], e["company"], e["bullets"])

    heading(main, "AI & Automation Projects")
    pe = profile.PROJECTS_ENTRY
    entry(main, pe["dates"], pe["loc"], pe["role"], pe["company"],
          cv.get("projects", profile.PROJECTS))

    heading(main, "Education")
    for ed in profile.EDUCATION:
        entry(main, ed["dates"], ed["loc"], ed["role"], ed["company"], [],
              company_desc=ed.get("desc"))

    # SIDE
    heading(side, "Summary", space_before=6, space_after=2)
    p = para(side, space_after=4)
    _run(p, cv.get("summary", profile.SUMMARY), size=7.8)

    heading(side, "Languages", space_before=6, space_after=2)
    side_list(side, profile.LANGUAGES)

    heading(side, "Tools & Technology", space_before=6, space_after=2)
    side_list(side, profile.TOOLS)

    heading(side, "Skills", space_before=6, space_after=2)
    side_list(side, [(None, s) for s in cv.get("skills", profile.SKILLS)])

    heading(side, "Core Focus", space_before=6, space_after=2)
    side_list(side, [(None, s) for s in cv.get("core_focus", profile.CORE_FOCUS)])

    doc.save(out_path)
    return out_path


# ============================================================
#  COVER LETTER
# ============================================================
def build_cover_letter(profile, job, out_path):
    cl = job.get("cover_letter", {})
    C = profile.CONTACT

    doc = Document()
    s = doc.sections[0]
    s.page_width = Inches(8.27); s.page_height = Inches(11.69)
    s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.5)
    s.left_margin = Inches(0.9); s.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = FONT; normal.font.size = Pt(10.5); normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(0); normal.paragraph_format.line_spacing = 1.12

    def para(space_before=0, space_after=6, align=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.12
        if align:
            p.alignment = align
        return p

    # header
    p = para(space_after=0); _run(p, C["name"], bold=True, size=19, color=NAVY)
    p = para(space_after=2)
    _run(p, cl.get("headline", C["headline"]), bold=True, size=8.8, color=NAVY)
    p = para(space_after=1)
    _run(p, f"{C['email']}   •   {C['phone']}   •   {C['location']}", size=8.6, color=MUTED)
    p = para(space_after=10)
    _run(p, f"{C['linkedin']}   •   {C['github']}", size=8.6, color=MUTED)
    add_bottom_rule(p, space="3")

    # date / addressee
    p = para(space_before=8, space_after=8); _run(p, cl.get("date", ""), size=10, color=MUTED)
    for line in cl.get("addressee_lines", []):
        p = para(space_after=0); _run(p, line, bold=(line is cl.get("addressee_lines", [None])[0]))
    if cl.get("addressee_lines"):
        # add breathing room before salutation
        doc.paragraphs[-1].paragraph_format.space_after = Pt(10)

    p = para(space_after=8); _run(p, cl.get("salutation", "Dear Hiring Team,"), size=10.5)

    for b in cl.get("paragraphs", []):
        p = para(space_after=8); _run(p, b, size=10.3)

    p = para(space_before=6, space_after=0); _run(p, "Warm regards,", size=10.3)
    p = para(space_before=2, space_after=0); _run(p, C["name"], bold=True, size=11, color=NAVY)

    doc.save(out_path)
    return out_path
